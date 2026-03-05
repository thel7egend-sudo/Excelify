import re

from docx import Document as DocxDocument
from docx.shared import Inches
from PySide6.QtCore import Qt, Signal
from PySide6.QtGui import (
    QColor,
    QFont,
    QFontDatabase,
    QFontMetrics,
    QKeyEvent,
    QMouseEvent,
    QTextCharFormat,
    QTextCursor,
    QTextDocument,
    QTextOption,
)
from PySide6.QtWidgets import (
    QComboBox,
    QFileDialog,
    QFrame,
    QGraphicsDropShadowEffect,
    QHBoxLayout,
    QLabel,
    QMessageBox,
    QPushButton,
    QScrollArea,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from spell_checker import SpellChecker


class PageTextEdit(QTextEdit):
    backspace_at_start = Signal()
    return_pressed = Signal()
    word_boundary_typed = Signal()

    _WORD_BOUNDARY_CHARS = set(" \t\n.,!?;:\"()[]{}")

    def __init__(self):
        super().__init__()
        self._spell_context_handler = None

    def keyPressEvent(self, event: QKeyEvent):
        if event.key() in (Qt.Key_Return, Qt.Key_Enter) and event.modifiers() == Qt.NoModifier:
            self.word_boundary_typed.emit()
            self.return_pressed.emit()
            return

        if (
            event.key() == Qt.Key_Backspace
            and not self.textCursor().hasSelection()
            and self.textCursor().position() == 0
        ):
            self.backspace_at_start.emit()
            return
        super().keyPressEvent(event)
        if event.modifiers() == Qt.NoModifier and event.text() in self._WORD_BOUNDARY_CHARS:
            self.word_boundary_typed.emit()

    def mousePressEvent(self, event: QMouseEvent):
        cursor = self.cursorForPosition(event.pos())
        doc_height = self.document().documentLayout().documentSize().height()
        if event.pos().y() >= doc_height:
            cursor.movePosition(QTextCursor.End)
        self.setTextCursor(cursor)
        super().mousePressEvent(event)

    def wheelEvent(self, event):
        # Keep each page fixed like a printed sheet: wheel scrolling should
        # scroll the outer document area, not create hidden scroll inside a page.
        event.ignore()

    def set_spell_context_handler(self, handler):
        self._spell_context_handler = handler

    def contextMenuEvent(self, event):
        menu = self.createStandardContextMenu()
        if self._spell_context_handler is not None:
            self._spell_context_handler(self, event, menu)
        menu.exec(event.globalPos())


class PageWidget(QFrame):
    PAGE_WIDTH = 800
    PAGE_HEIGHT = 1100
    PAGE_MARGIN = 56

    def __init__(self):
        super().__init__()
        self.setObjectName("docPage")
        self.setFixedSize(self.PAGE_WIDTH, self.PAGE_HEIGHT)
        self.setFrameShape(QFrame.NoFrame)

        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(28)
        shadow.setXOffset(0)
        shadow.setYOffset(4)
        shadow.setColor(QColor(0, 0, 0, 45))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(self.PAGE_MARGIN, self.PAGE_MARGIN, self.PAGE_MARGIN, self.PAGE_MARGIN)
        layout.setSpacing(0)

        self.editor = PageTextEdit()
        self.editor.setFrameShape(QFrame.NoFrame)
        self.editor.setAcceptRichText(False)
        self.editor.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.editor.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.editor.setLineWrapMode(QTextEdit.WidgetWidth)
        self.editor.setFixedHeight(self.usable_page_height)
        layout.addWidget(self.editor)

    @property
    def usable_page_width(self):
        return self.PAGE_WIDTH - (self.PAGE_MARGIN * 2)

    @property
    def usable_page_height(self):
        return self.PAGE_HEIGHT - (self.PAGE_MARGIN * 2)


class WordStyleEditor(QWidget):
    textChanged = Signal()

    PAGE_GAP = 48
    THEME_TOKENS = {
        "theme-light": {
            "workspace": "#dfe1e5",
            "page_text": "#111827",
            "page_bg": "#ffffff",
            "page_border": "#d8dde6",
        },
        "theme-dark": {
            "workspace": "#13161c",
            "page_text": "#eaeaea",
            "page_bg": "#1d1f24",
            "page_border": "rgba(255,255,255,0.05)",
        },
    }

    def __init__(self, initial_text=""):
        super().__init__()
        self._is_reflowing = False
        self._active_page_idx = 0
        self._pages = []
        self._word_regex = re.compile(r"[A-Za-z']+")
        self._block_spelling_ranges = {}
        self._page_text_lengths = {}
        self._spell_checker = SpellChecker()

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.scroll_area.setFrameShape(QFrame.NoFrame)

        self.container = QWidget()
        self.pages_layout = QVBoxLayout(self.container)
        self.pages_layout.setContentsMargins(24, 24, 24, 24)
        self.pages_layout.setSpacing(self.PAGE_GAP)
        self.pages_layout.setAlignment(Qt.AlignHCenter | Qt.AlignTop)

        self.scroll_area.setWidget(self.container)
        root.addWidget(self.scroll_area)

        self._apply_theme("theme-light")
        self.setPlainText(initial_text)

    @property
    def usable_page_width(self):
        return PageWidget.PAGE_WIDTH - (PageWidget.PAGE_MARGIN * 2)

    @property
    def usable_page_height(self):
        return PageWidget.PAGE_HEIGHT - (PageWidget.PAGE_MARGIN * 2)

    def _create_page(self, text=""):
        page = PageWidget()
        page.editor.document().setDocumentMargin(0)
        text_option = page.editor.document().defaultTextOption()
        text_option.setWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        page.editor.document().setDefaultTextOption(text_option)
        page.editor.setPlainText(text)
        page.editor.set_spell_context_handler(self._show_spell_context_menu)
        self._page_text_lengths[page.editor] = len(text)
        page.editor.textChanged.connect(lambda p=page: self._on_page_text_changed(p))
        page.editor.backspace_at_start.connect(lambda p=page: self._merge_with_previous(p))
        page.editor.return_pressed.connect(lambda p=page: self._handle_return_pressed(p))
        page.editor.word_boundary_typed.connect(lambda p=page: self._check_word_before_cursor(p))
        page.editor.selectionChanged.connect(lambda p=page: self._track_active_page(p))
        page.editor.cursorPositionChanged.connect(lambda p=page: self._track_active_page(p))
        page.editor.cursorPositionChanged.connect(lambda p=page: self._ensure_page_cursor_visible(p))
        page.editor.verticalScrollBar().rangeChanged.connect(lambda *_args, e=page.editor: e.verticalScrollBar().setValue(0))
        return page

    def _word_bounds_at_position(self, text, pos):
        if not text:
            return None
        idx = min(max(0, pos), len(text) - 1)
        if not re.match(r"[A-Za-z']", text[idx]):
            if idx > 0 and re.match(r"[A-Za-z']", text[idx - 1]):
                idx -= 1
            else:
                return None

        start = idx
        while start > 0 and re.match(r"[A-Za-z']", text[start - 1]):
            start -= 1
        end = idx + 1
        while end < len(text) and re.match(r"[A-Za-z']", text[end]):
            end += 1
        return start, end

    def _check_word_before_cursor(self, page):
        if page not in self._pages:
            return

        editor = page.editor
        cursor = editor.textCursor()
        pos = cursor.position()
        text = editor.toPlainText()
        if pos <= 0 or pos > len(text):
            return

        bounds = self._word_bounds_at_position(text, pos - 1)
        if not bounds:
            return

        word_start, word_end = bounds
        word = text[word_start:word_end]
        if not word:
            return

        block = cursor.block()
        block_number = block.blockNumber()
        editor_ranges = self._block_spelling_ranges.setdefault(editor, {})
        ranges = editor_ranges.setdefault(block_number, [])
        ranges = [r for r in ranges if not (r[0] == word_start and r[1] == word_end)]
        if not self._spell_checker.is_correct(word):
            ranges.append((word_start, word_end))
        if ranges:
            editor_ranges[block_number] = ranges
        else:
            editor_ranges.pop(block_number, None)
        self._refresh_spell_underlines(editor)

    def _full_spell_check_page(self, page):
        if page not in self._pages:
            return

        editor = page.editor
        block_ranges = {}
        block = editor.document().firstBlock()
        while block.isValid():
            local_ranges = []
            for match in self._word_regex.finditer(block.text()):
                word = match.group(0)
                if not self._spell_checker.is_correct(word):
                    start = block.position() + match.start()
                    end = block.position() + match.end()
                    local_ranges.append((start, end))
            if local_ranges:
                block_ranges[block.blockNumber()] = local_ranges
            block = block.next()

        self._block_spelling_ranges[editor] = block_ranges
        self._refresh_spell_underlines(editor)

    def _spell_check_current_block(self, page):
        if page not in self._pages:
            return

        editor = page.editor
        cursor = editor.textCursor()
        block = cursor.block()
        block_number = block.blockNumber()
        block_start = block.position()

        local_ranges = []
        for match in self._word_regex.finditer(block.text()):
            word = match.group(0)
            if not self._spell_checker.is_correct(word):
                local_ranges.append((block_start + match.start(), block_start + match.end()))

        editor_ranges = self._block_spelling_ranges.setdefault(editor, {})
        if local_ranges:
            editor_ranges[block_number] = local_ranges
        else:
            editor_ranges.pop(block_number, None)
        self._refresh_spell_underlines(editor)

    def _refresh_spell_underlines(self, editor):
        selections = []
        fmt = QTextCharFormat()
        fmt.setUnderlineStyle(QTextCharFormat.SpellCheckUnderline)
        fmt.setUnderlineColor(Qt.red)
        for ranges in self._block_spelling_ranges.get(editor, {}).values():
            for start, end in ranges:
                cursor = editor.textCursor()
                cursor.setPosition(start)
                cursor.setPosition(end, QTextCursor.KeepAnchor)
                sel = QTextEdit.ExtraSelection()
                sel.cursor = cursor
                sel.format = fmt
                selections.append(sel)
        editor.setExtraSelections(selections)

    def _replace_word_in_editor(self, editor, start, end, replacement):
        cursor = editor.textCursor()
        cursor.setPosition(start)
        cursor.setPosition(end, QTextCursor.KeepAnchor)
        cursor.insertText(replacement)
        editor.setTextCursor(cursor)

    def _show_spell_context_menu(self, editor, event, menu):
        click_cursor = editor.cursorForPosition(event.pos())
        text = editor.toPlainText()
        bounds = self._word_bounds_at_position(text, click_cursor.position())
        if not bounds:
            return

        word_start, word_end = bounds
        word = text[word_start:word_end]
        if self._spell_checker.is_correct(word):
            return

        suggestions = self._spell_checker.suggest(word, limit=5)
        if suggestions:
            for suggestion in suggestions[::-1]:
                action = menu.addAction(suggestion)
                action.triggered.connect(
                    lambda _checked=False, s=suggestion, st=word_start, en=word_end, e=editor: self._replace_word_in_editor(e, st, en, s)
                )

        ignore_action = menu.addAction("(ignore)")
        ignore_action.triggered.connect(lambda: None)
        menu.addSeparator()

    def _ensure_page_cursor_visible(self, page):
        if page not in self._pages:
            return
        cursor_center = page.editor.mapTo(self.container, page.editor.cursorRect().center())
        self.scroll_area.ensureVisible(cursor_center.x(), cursor_center.y(), 24, 48)

    def _append_page(self, text=""):
        page = self._create_page(text)
        self._pages.append(page)
        self.pages_layout.addWidget(page)
        return page

    def _insert_page_after(self, idx, text=""):
        page = self._create_page(text)
        self._pages.insert(idx + 1, page)
        self.pages_layout.insertWidget(idx + 1, page)
        return page

    def _remove_page(self, idx):
        if len(self._pages) <= 1:
            return
        page = self._pages.pop(idx)
        self._block_spelling_ranges.pop(page.editor, None)
        self._page_text_lengths.pop(page.editor, None)
        self.pages_layout.removeWidget(page)
        page.deleteLater()

    def _track_active_page(self, page):
        if page in self._pages:
            self._active_page_idx = self._pages.index(page)

    def _text_fits(self, text, editor=None):
        if not self._pages:
            return True

        source_editor = editor or self._pages[0].editor
        probe = QTextDocument()
        probe.setDefaultFont(source_editor.font())
        probe.setDocumentMargin(0)
        probe.setDefaultTextOption(source_editor.document().defaultTextOption())
        probe.setTextWidth(self.usable_page_width)
        probe.setPlainText(text)
        doc_height = probe.documentLayout().documentSize().height()

        line_spacing = QFontMetrics(source_editor.font()).lineSpacing()
        max_height = max(0.0, self.usable_page_height - line_spacing)
        return doc_height <= max_height

    def _handle_return_pressed(self, page):
        if self._is_reflowing or page not in self._pages:
            return

        editor = page.editor
        cursor = editor.textCursor()
        if cursor.hasSelection():
            cursor.removeSelectedText()

        idx = self._pages.index(page)
        text = editor.toPlainText()
        pos = cursor.position()
        candidate = f"{text[:pos]}\n{text[pos:]}"

        self._is_reflowing = True
        if self._text_fits(candidate, editor):
            cursor.insertBlock()
            editor.setTextCursor(cursor)
            self._ensure_page_cursor_visible(page)
            self._is_reflowing = False
            self.textChanged.emit()
            return

        if idx + 1 >= len(self._pages):
            self._append_page("")

        next_editor = self._pages[idx + 1].editor
        self._is_reflowing = False

        next_cursor = next_editor.textCursor()
        next_cursor.setPosition(0)
        next_editor.setTextCursor(next_cursor)
        next_editor.setFocus()
        self._ensure_page_cursor_visible(self._pages[idx + 1])
        self._active_page_idx = idx + 1
        self.textChanged.emit()

    def _fitting_index(self, text):
        if not text:
            return 0
        low, high = 0, len(text)
        fit = 0
        while low <= high:
            mid = (low + high) // 2
            chunk = text[:mid]
            if self._text_fits(chunk):
                fit = mid
                low = mid + 1
            else:
                high = mid - 1

        split_at = fit
        while 1 < split_at < len(text) and not text[split_at - 1].isspace() and text[split_at].isalnum():
            split_at -= 1

        return max(1, split_at)

    def _global_cursor_position(self, page_idx, local_pos):
        clamped_page_idx = min(max(0, page_idx), len(self._pages) - 1)
        prior_text_len = sum(len(self._pages[i].editor.toPlainText()) for i in range(clamped_page_idx))
        current_len = len(self._pages[clamped_page_idx].editor.toPlainText())
        return prior_text_len + max(0, min(local_pos, current_len))

    def _position_from_global_offset(self, global_offset):
        if not self._pages:
            return 0, 0

        remaining = max(0, global_offset)
        for idx, page in enumerate(self._pages):
            text_len = len(page.editor.toPlainText())
            if remaining <= text_len or idx == len(self._pages) - 1:
                return idx, min(remaining, text_len)
            remaining -= text_len

        last_idx = len(self._pages) - 1
        return last_idx, len(self._pages[last_idx].editor.toPlainText())

    def _capture_caret_state(self):
        if not self._pages:
            return None

        focus_widget = self.focusWidget()
        if isinstance(focus_widget, PageTextEdit) and focus_widget.parent() in self._pages:
            editor = focus_widget
            page_idx = self._pages.index(focus_widget.parent())
        else:
            page_idx = min(max(0, self._active_page_idx), len(self._pages) - 1)
            editor = self._pages[page_idx].editor

        cursor = editor.textCursor()
        return {
            "anchor": self._global_cursor_position(page_idx, cursor.anchor()),
            "position": self._global_cursor_position(page_idx, cursor.position()),
            "had_focus": editor.hasFocus(),
        }

    def _restore_caret_state(self, caret_state):
        if not caret_state or not self._pages:
            return

        anchor_idx, anchor_local = self._position_from_global_offset(caret_state["anchor"])
        pos_idx, pos_local = self._position_from_global_offset(caret_state["position"])

        target_editor = self._pages[pos_idx].editor
        cursor = target_editor.textCursor()

        if anchor_idx == pos_idx:
            cursor.setPosition(anchor_local)
            if anchor_local != pos_local:
                cursor.setPosition(pos_local, QTextCursor.KeepAnchor)
            else:
                cursor.setPosition(pos_local)
        else:
            cursor.setPosition(pos_local)

        target_editor.setTextCursor(cursor)
        target_editor.ensureCursorVisible()
        self._ensure_page_cursor_visible(self._pages[pos_idx])
        if caret_state.get("had_focus", True):
            target_editor.setFocus()
        self._active_page_idx = pos_idx

    def _rebalance_from(self, start_idx):
        idx = max(0, start_idx)
        while idx < len(self._pages):
            current = self._pages[idx].editor
            text = current.toPlainText()
            if self._text_fits(text):
                idx += 1
                continue

            split_at = self._fitting_index(text)
            leading = text[:split_at]
            trailing = text[split_at:]

            current.blockSignals(True)
            current.setPlainText(leading)
            current.blockSignals(False)

            if idx + 1 >= len(self._pages):
                self._append_page("")

            next_editor = self._pages[idx + 1].editor
            next_text = next_editor.toPlainText()
            next_editor.blockSignals(True)
            next_editor.setPlainText(trailing + next_text)
            next_editor.blockSignals(False)

        self._pull_text_up(max(0, start_idx - 1))

    def _pull_text_up(self, start_idx):
        idx = max(0, start_idx)
        while idx < len(self._pages) - 1:
            current = self._pages[idx].editor
            nxt = self._pages[idx + 1].editor
            current_text = current.toPlainText()
            next_text = nxt.toPlainText()

            if not next_text:
                self._remove_page(idx + 1)
                continue

            if not self._text_fits(current_text + next_text[:1]):
                idx += 1
                continue

            low, high, fit = 1, len(next_text), 0
            while low <= high:
                mid = (low + high) // 2
                candidate = current_text + next_text[:mid]
                if self._text_fits(candidate):
                    fit = mid
                    low = mid + 1
                else:
                    high = mid - 1

            moved = next_text[:fit]
            remainder = next_text[fit:]

            current.blockSignals(True)
            current.setPlainText(current_text + moved)
            current.blockSignals(False)

            nxt.blockSignals(True)
            nxt.setPlainText(remainder)
            nxt.blockSignals(False)

            if not remainder:
                self._remove_page(idx + 1)
            else:
                idx += 1

    def _merge_with_previous(self, page):
        if page not in self._pages:
            return
        idx = self._pages.index(page)
        if idx == 0:
            return

        prev_editor = self._pages[idx - 1].editor
        this_editor = self._pages[idx].editor

        prev_text = prev_editor.toPlainText()
        this_text = this_editor.toPlainText()
        caret_state = self._capture_caret_state()
        if caret_state:
            caret_state["position"] = self._global_cursor_position(idx - 1, len(prev_text))
            caret_state["anchor"] = caret_state["position"]

        self._is_reflowing = True
        prev_editor.blockSignals(True)
        prev_editor.setPlainText(prev_text + this_text)
        prev_editor.blockSignals(False)
        self._remove_page(idx)
        self._rebalance_from(idx - 1)
        self._is_reflowing = False
        self._restore_caret_state(caret_state)

        self.textChanged.emit()

    def _on_page_text_changed(self, page):
        if self._is_reflowing:
            return
        if page not in self._pages:
            return

        self._is_reflowing = True
        caret_state = self._capture_caret_state()
        idx = self._pages.index(page)
        self._rebalance_from(idx)
        self._is_reflowing = False
        self._restore_caret_state(caret_state)
        for pg in self._pages:
            pg.editor.verticalScrollBar().setValue(0)

        current_len = len(page.editor.toPlainText())
        previous_len = self._page_text_lengths.get(page.editor, current_len)
        if abs(current_len - previous_len) > 120:
            self._full_spell_check_page(page)
        else:
            self._spell_check_current_block(page)
        self._page_text_lengths[page.editor] = current_len

        self.textChanged.emit()

    def toPlainText(self):
        return "".join(page.editor.toPlainText() for page in self._pages)

    def setPlainText(self, text):
        self._is_reflowing = True
        for page in self._pages:
            self.pages_layout.removeWidget(page)
            page.deleteLater()
        self._pages = []
        self._block_spelling_ranges = {}
        self._page_text_lengths = {}

        self._append_page(text or "")
        self._rebalance_from(0)
        self._is_reflowing = False

        for page in self._pages:
            self._full_spell_check_page(page)

        if self._pages:
            self._pages[0].editor.moveCursor(QTextCursor.Start)
            self._pages[0].editor.setFocus()

    def _apply_theme(self, theme_class: str):
        tokens = self.THEME_TOKENS[theme_class]
        self.setProperty("theme", theme_class)
        self.container.setProperty("theme", theme_class)
        self.style().unpolish(self)
        self.style().polish(self)
        self.style().unpolish(self.container)
        self.style().polish(self.container)

        self.setStyleSheet(
            """
            QWidget[theme="theme-light"],
            QWidget[theme="theme-dark"] {
                background: %(workspace)s;
            }
            QScrollArea {
                border: none;
                background: %(workspace)s;
            }
            QFrame#docPage {
                background: %(page_bg)s;
                border: 1px solid %(page_border)s;
            }
            QFrame#docPage QTextEdit {
                background: transparent;
                color: %(page_text)s;
                border: none;
                font-size: 14px;
            }
            """
            % tokens
        )

    def set_dark_mode(self, enabled: bool):
        self._apply_theme("theme-dark" if enabled else "theme-light")


class DocEditorPage(QWidget):
    document_changed = Signal()
    export_requested = Signal(object)

    def __init__(self, document):
        super().__init__()
        self.document = document
        self.setProperty("theme", "theme-light")
        self._font_categories = {
            "Sans Serif": ["Segoe UI", "Arial", "Calibri", "Verdana", "Tahoma"],
            "Serif": ["Times New Roman", "Georgia", "Cambria", "Garamond"],
            "Monospace": ["Consolas", "Courier New", "Lucida Console"],
            "Creative": ["Comic Sans MS", "Impact", "Trebuchet MS"],
        }
        self._active_font_family = "Segoe UI"
        self._active_font_size = 14.0

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self.ribbon = QWidget()
        self.ribbon.setObjectName("docRibbon")
        self.ribbon.setFixedHeight(72)
        ribbon_layout = QHBoxLayout(self.ribbon)
        ribbon_layout.setContentsMargins(16, 0, 16, 0)

        placeholder_btn = QPushButton("Tools (Coming Soon)")
        placeholder_btn.setEnabled(False)
        placeholder_btn.setFixedHeight(32)
        ribbon_layout.addWidget(placeholder_btn)

        font_family_label = QLabel("Font")
        ribbon_layout.addWidget(font_family_label)

        self.font_family_combo = QComboBox()
        self.font_family_combo.setFixedHeight(32)
        self._populate_font_families()
        self.font_family_combo.currentTextChanged.connect(self._apply_font_family)
        ribbon_layout.addWidget(self.font_family_combo)

        font_size_label = QLabel("Font Size")
        ribbon_layout.addWidget(font_size_label)

        self.font_size_combo = QComboBox()
        self.font_size_combo.addItems(["10", "11", "12", "14", "16", "18", "20", "24"])
        self.font_size_combo.setCurrentText("14")
        self.font_size_combo.setFixedHeight(32)
        self.font_size_combo.currentTextChanged.connect(self._apply_font_size)
        ribbon_layout.addWidget(self.font_size_combo)

        ribbon_layout.addStretch()

        self.export_btn = QPushButton("Export to Docs")
        self.export_btn.setFixedHeight(36)
        self.export_btn.clicked.connect(lambda: self.export_requested.emit(self.document))
        ribbon_layout.addWidget(self.export_btn)

        self.editor = WordStyleEditor(self.document.content or "")
        self.editor.textChanged.connect(self._on_text_changed)

        layout.addWidget(self.ribbon)
        layout.addWidget(self.editor)

        self._apply_font_family(self.font_family_combo.currentText())
        self._apply_font_size(self.font_size_combo.currentText())

        self.apply_grid_dark_mode(False)

    @property
    def _usable_page_width(self):
        return self.editor.usable_page_width

    @property
    def _usable_page_height(self):
        return self.editor.usable_page_height

    def _on_text_changed(self):
        self.document.content = self.editor.toPlainText()
        self.document_changed.emit()

    def _apply_font_size(self, size_text):
        if not size_text:
            return
        size = float(size_text)
        self._active_font_size = size
        if not self.editor._pages:
            return
        self._apply_editor_font_settings()

    def _populate_font_families(self):
        db = QFontDatabase()
        available = set(db.families())

        for category, families in self._font_categories.items():
            for family in families:
                if family not in available:
                    continue
                self.font_family_combo.addItem(f"{category} — {family}", family)

        if self.font_family_combo.count() == 0:
            self.font_family_combo.addItem("Default — Sans Serif", self._active_font_family)

        current_idx = self.font_family_combo.findData(self._active_font_family)
        self.font_family_combo.setCurrentIndex(max(0, current_idx))

    def _apply_font_family(self, display_text):
        del display_text
        selected_family = self.font_family_combo.currentData()
        if not selected_family:
            return
        self._active_font_family = selected_family
        self._apply_editor_font_settings()

    def _apply_editor_font_settings(self):
        if not self.editor._pages:
            return

        self.editor._is_reflowing = True
        try:
            caret_state = self.editor._capture_caret_state()
            for page in self.editor._pages:
                page_font = QFont(self._active_font_family)
                page_font.setPointSizeF(self._active_font_size)
                page.editor.setFont(page_font)
                page.editor.document().setDefaultFont(page_font)
            self.editor._rebalance_from(0)
            self.editor._restore_caret_state(caret_state)
        finally:
            self.editor._is_reflowing = False

    def _paginate_text(self, text):
        if not text:
            return [""]

        probe = QTextDocument()
        probe.setDocumentMargin(0)
        probe.setDefaultFont(self.editor.font())
        probe.setTextWidth(self._usable_page_width)
        max_height = self._usable_page_height
        pages = []
        remaining = text

        while remaining:
            probe.setPlainText(remaining)
            if probe.documentLayout().documentSize().height() <= max_height:
                pages.append(remaining)
                break

            low = 1
            high = len(remaining)
            fit = 1

            while low <= high:
                mid = (low + high) // 2
                chunk = remaining[:mid]
                probe.setPlainText(chunk)
                if probe.documentLayout().documentSize().height() <= max_height:
                    fit = mid
                    low = mid + 1
                else:
                    high = mid - 1

            split_at = fit
            while split_at > 1 and not remaining[split_at - 1].isspace():
                split_at -= 1

            if split_at <= 1:
                split_at = fit

            pages.append(remaining[:split_at])
            remaining = remaining[split_at:]

        return pages or [""]

    def export_to_docx(self):
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Export to Docs",
            f"{self.document.name}.docx",
            "Word Documents (*.docx)",
        )

        if not path:
            return

        if not path.lower().endswith(".docx"):
            path = f"{path}.docx"

        try:
            text = self.editor.toPlainText()
            page_chunks = self._paginate_text(text)

            doc = DocxDocument()
            section = doc.sections[0]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)

            for page_index, chunk in enumerate(page_chunks):
                paragraphs = chunk.split("\n")
                for paragraph in paragraphs:
                    doc.add_paragraph(paragraph)

                if page_index < len(page_chunks) - 1:
                    doc.add_page_break()

            doc.save(path)
        except Exception as e:
            QMessageBox.critical(self, "Export Failed", f"Could not export file:\n{e}")
            return

        QMessageBox.information(self, "Export Complete", "Word document exported successfully.")

    def apply_grid_dark_mode(self, enabled: bool):
        theme_class = "theme-dark" if enabled else "theme-light"
        self.setProperty("theme", theme_class)
        self.style().unpolish(self)
        self.style().polish(self)
        self.editor.set_dark_mode(enabled)

        theme_tokens = {
            "theme-light": {
                "bg": "#f7f8fa",
                "text": "#111827",
                "ribbon_bg": "#f9fafb",
                "ribbon_border": "#e5e7eb",
            },
            "theme-dark": {
                "bg": "#202124",
                "text": "#eaeaea",
                "ribbon_bg": "#252525",
                "ribbon_border": "rgba(255, 255, 255, 0.06)",
            },
        }[theme_class]

        self.setStyleSheet(
            """
            QWidget {
                background: %(bg)s;
                color: %(text)s;
            }
            QWidget#docRibbon {
                background-color: %(ribbon_bg)s;
                border-bottom: 1px solid %(ribbon_border)s;
            }
            """
            % theme_tokens
        )
